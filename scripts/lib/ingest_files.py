"""Extract immutable local career sources and register them in the manifest."""

from __future__ import annotations

import argparse
from pathlib import Path
from typing import Any

from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader

from scripts.lib.source_manifest import (
    find_source_by_hash,
    load_manifest,
    register_source,
    repository_path,
    sha256_file,
    source_id,
)

REPOSITORY_ROOT = Path(__file__).resolve().parents[2]
SOURCES_ROOT = REPOSITORY_ROOT / "sources"
MANIFEST_PATH = SOURCES_ROOT / "manifest.yaml"
EXTRACTED_ROOT = SOURCES_ROOT / "extracted"

SOURCE_TYPES = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".md": "markdown",
    ".markdown": "markdown",
    ".txt": "text",
    ".html": "html",
    ".htm": "html",
    ".png": "image",
    ".jpg": "image",
    ".jpeg": "image",
    ".webp": "image",
}


class IngestionError(ValueError):
    """Raised when a source cannot be safely ingested."""


def ensure_source_path(path: Path, sources_root: Path) -> Path:
    resolved_path = path.resolve()
    resolved_sources = sources_root.resolve()
    try:
        relative_path = resolved_path.relative_to(resolved_sources)
    except ValueError as error:
        raise IngestionError(f"source must be inside {sources_root}") from error
    if relative_path.parts and relative_path.parts[0] == "extracted":
        raise IngestionError("generated extracted text cannot be ingested as an original")
    if not resolved_path.is_file():
        raise IngestionError(f"source file does not exist: {path}")
    return resolved_path


def extract_text(path: Path, source_type: str) -> tuple[str | None, dict[str, Any]]:
    if source_type == "pdf":
        reader = PdfReader(str(path))
        text = "\n\n".join((page.extract_text() or "").strip() for page in reader.pages)
        return text.strip(), {"pages": len(reader.pages)}

    if source_type == "docx":
        document = Document(str(path))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
        return "\n".join(value for value in paragraphs if value), {
            "paragraphs": len(document.paragraphs)
        }

    if source_type == "html":
        raw_html = path.read_text(encoding="utf-8", errors="replace")
        soup = BeautifulSoup(raw_html, "html.parser")
        title = str(soup.title.string) if soup.title and soup.title.string else None
        return soup.get_text("\n", strip=True), {"title": title}

    if source_type in {"markdown", "text"}:
        return path.read_text(encoding="utf-8", errors="replace"), {}

    if source_type == "image":
        return None, {}

    raise IngestionError(f"unsupported source type: {source_type}")


def ingest_file(
    path: Path,
    *,
    repository_root: Path = REPOSITORY_ROOT,
    sources_root: Path | None = None,
    manifest_path: Path | None = None,
    extracted_root: Path | None = None,
) -> tuple[dict[str, Any], bool]:
    sources_root = sources_root or repository_root / "sources"
    manifest_path = manifest_path or sources_root / "manifest.yaml"
    extracted_root = extracted_root or sources_root / "extracted"
    source_path = ensure_source_path(path, sources_root)

    source_type = SOURCE_TYPES.get(source_path.suffix.lower())
    if source_type is None:
        raise IngestionError(f"unsupported source extension: {source_path.suffix}")

    digest_before = sha256_file(source_path)
    existing = find_source_by_hash(load_manifest(manifest_path), digest_before, source_type)
    if existing is not None:
        return existing, False

    text, extraction_metadata = extract_text(source_path, source_type)
    digest_after = sha256_file(source_path)
    if digest_before != digest_after:
        raise IngestionError(f"source changed during ingestion: {source_path}")

    entry_id = source_id(source_path.stem, digest_before)
    extracted_text_path: str | None = None
    if text is not None:
        extracted_root.mkdir(parents=True, exist_ok=True)
        text_path = extracted_root / f"{entry_id}.txt"
        text_path.write_text(text, encoding="utf-8")
        extracted_text_path = repository_path(text_path, repository_root)

    relative_source_path = repository_path(source_path, repository_root)
    metadata = {
        "filename": source_path.name,
        "size_bytes": source_path.stat().st_size,
        **extraction_metadata,
    }
    return register_source(
        manifest_path,
        entry_id=entry_id,
        source_type=source_type,
        path=relative_source_path,
        origin=relative_source_path,
        digest=digest_before,
        extracted_text_path=extracted_text_path,
        metadata=metadata,
    )


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Ingest immutable files under sources/.")
    parser.add_argument("paths", nargs="+", type=Path)
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    for path in args.paths:
        entry, created = ingest_file(path)
        action = "ingested" if created else "already registered"
        print(f"{action} {entry['id']}: {entry['path']}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
